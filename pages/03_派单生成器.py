import os

import streamlit as st
import pandas as pd
from datetime import date
from PIL import Image
from docx import Document
from docx.shared import Inches

from home import df_survey, df_treatment, df_first_damage


# 定义日期选择器函数
def show_select_data():
    selected_date = st.date_input('选择一个日期', value=date.today())
    # 用选定的日期来过滤数据
    filtered_data = df_survey[df_survey['调查日期'] == pd.to_datetime(selected_date)]
    filtered_data.reset_index(inplace=True, drop=True)
    # 展示过滤后的数据
    st.subheader('当天数据')
    st.dataframe(filtered_data, use_container_width=True)
    return filtered_data


# 定义用于获取上传图片的函数
def get_images():
    uploaded_files = st.file_uploader("请选择图片", accept_multiple_files=True, type=['png', 'jpg', 'jpeg'])
    if uploaded_files is not None:
        return uploaded_files


def process_word(row, template_file, img_files):
    # 解析每一行数据
    date = row['调查日期']  # 日期
    order_number = row['序号']  # 序号
    town = str(row['乡镇/街道'])  # 乡镇
    number = str(row['点位编号'])  # 编号
    location = str(row['点位名'])  # 点位名
    report_time = str(row['上报时间'])  # 上报时间
    reporter = "李继广"  # 上报人,默认为李继广
    pest = "美国白蛾"  # 虫种,默认为美国白蛾
    host = str(row['危害寄主'])  # 寄主
    number_of_trees = str(row['受害株数'])  # 棵树
    greenery_type = str(row['地块类型'])  # 绿化性质
    description = str(row['详细描述'])  # 详细描述
    webbed_nests = str(row['网幕数'])  # 网幕数

    doc = Document(template_file)  # 读取Word模板文件
    # 替换第一段落的内容
    text = '                                                    编号：'
    if doc.paragraphs[0].text == text:
        doc.paragraphs[0].text = text + '{}-{}'.format(date, order_number)
    # 遍历Word文档中的表格，修改表格中的内容
    for table in doc.tables:
        # 在表格中查找占位符并进行替换
        for row in table.rows:
            for cell in row.cells:
                # 修改表格中的具体街道
                if cell.text == '具体街道':
                    cell.text = town
                    cell.paragraphs[0].runs[0].bold = True  # 设置字体为粗体
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体点位
                if cell.text == '具体点位':
                    cell.text = number
                    cell.paragraphs[0].runs[0].bold = True  # 设置字体为粗体
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体位置
                if cell.text == '具体位置':
                    cell.text = location
                    cell.paragraphs[0].runs[0].bold = True  # 设置字体为粗体
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体上报时间
                if cell.text == '具体上报时间':
                    cell.text = report_time
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体上报人
                if cell.text == '具体上报人':
                    cell.text = reporter
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体虫种
                if cell.text == '具体虫种':
                    cell.text = pest
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体寄主
                if cell.text == '具体寄主':
                    cell.text = host
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体株数
                if cell.text == '具体株数':
                    cell.text = number_of_trees
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体绿地性质
                if cell.text == '具体绿化性质':
                    cell.text = greenery_type
                    cell.paragraphs[0].paragraph_format.alignment = 1  # 居中
                # 修改表格中的具体描述
                if cell.text == '具体描述':
                    cell.text = description
                # 修改表格中的具体网幕数
                if cell.text == '具体网幕数':
                    cell.text = webbed_nests
                    cell.paragraphs[0].paragraph_format.alignment = 1
    # 插入图片
    # 首先筛选出前缀为编号的图片
    img_files = [file for file in img_files if file.name.startswith(number)]
    count = 0
    # 4个图片的索引
    size_list = [(9, 0), (9, 3), (11, 0), (11, 3)]
    for img in img_files:
        image = Image.open(img)
        width, height = image.size
        row, col = size_list[count]
        if width > height:
            w = Inches(3)
            h = Inches(2)
        else:
            w = Inches(2)
            h = Inches(3)
        table = doc.tables[0]
        cell = table.cell(row, col)
        cell.paragraphs[0].add_run().add_picture(img, width=w, height=h)
        count += 1

    # 保存Word文档
    doc.save('派单/2023林业有害生物防治工作单({0})-{1}-{2}-{3}.doc'.format(town, location, date, order_number))
    return None


# 定义一个用于生成派单的函数的按钮
def generate_word(filtered_data, template_file, img_files):
    # 判断是否存在「派单」文件夹
    if not os.path.exists("派单"):
        # 创建「派单」文件夹
        os.mkdir("派单")
    # 按钮
    button = st.button('生成派单')
    if button:
        # 对数据筛选不合格的数据
        filtered_data = filtered_data[filtered_data['调查结果'] == '不合格']
        # 为筛选后的数据添加一列序号
        # filtered_data['序号'] = range(1, len(filtered_data) + 1)
        filtered_data.loc[:, '序号'] = range(1, len(filtered_data) + 1)
        filtered_data['调查日期'] = filtered_data['调查日期'].dt.strftime('%Y-%m-%d')

        template_file = template_file
        img_files = img_files
        # 对筛选后的数据进行处理
        if filtered_data is not None:
            # 遍历筛选出的数据
            for index, row in filtered_data.iterrows():
                # 生成Word文件
                process_word(row=row, template_file=template_file, img_files=img_files)
        # 生成完成后，弹出提示
        st.success('派单生成完成！')
        st.success('请在「派单」文件夹中查看！')
        st.success('共生成{}份派单！'.format(len(filtered_data)))
    else:
        st.info('请点击生成派单按钮！')


if __name__ == '__main__':
    st.title('派单生成器')
    filtered_data = show_select_data()
    template_file = r'template/美国白蛾派单模版.docx'
    img_files = get_images()
    generate_word(filtered_data, template_file, img_files)
